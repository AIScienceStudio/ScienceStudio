#!/usr/bin/env python3
"""
DOCX MCP Server for ScienceStudio

Provides tools for Claude to read and write Word documents:
- Read document content
- Write/create documents
- Insert text at positions
- Handle comments and track changes
"""

import asyncio
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent

# Initialize MCP server
server = Server("docx-mcp")


def read_docx(path: str) -> dict:
    """Read content from a Word document."""
    doc = Document(path)

    content = {
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": []
    }

    # Extract paragraphs with style info
    for para in doc.paragraphs:
        para_data = {
            "text": para.text,
            "style": para.style.name if para.style else None,
            "alignment": str(para.alignment) if para.alignment else None
        }
        content["paragraphs"].append(para_data)

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)

    # Extract headers
    for section in doc.sections:
        if section.header:
            header_text = "\n".join([p.text for p in section.header.paragraphs])
            if header_text.strip():
                content["headers"].append(header_text)

        if section.footer:
            footer_text = "\n".join([p.text for p in section.footer.paragraphs])
            if footer_text.strip():
                content["footers"].append(footer_text)

    return content


def read_docx_plain(path: str) -> str:
    """Read document as plain text."""
    doc = Document(path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    return "\n\n".join(text_parts)


def write_docx(path: str, content: str, title: str = None) -> dict:
    """Create a new Word document with content."""
    doc = Document()

    # Add title if provided
    if title:
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add content paragraphs
    paragraphs = content.split("\n\n")
    for para_text in paragraphs:
        if para_text.strip():
            # Check if it's a heading (starts with #)
            if para_text.startswith("# "):
                doc.add_heading(para_text[2:], 1)
            elif para_text.startswith("## "):
                doc.add_heading(para_text[3:], 2)
            elif para_text.startswith("### "):
                doc.add_heading(para_text[4:], 3)
            else:
                doc.add_paragraph(para_text)

    # Save document
    doc.save(path)

    return {
        "success": True,
        "path": path,
        "paragraphs": len(doc.paragraphs)
    }


def append_to_docx(path: str, content: str) -> dict:
    """Append content to an existing document."""
    if not Path(path).exists():
        return {"success": False, "error": f"File not found: {path}"}

    doc = Document(path)

    paragraphs = content.split("\n\n")
    for para_text in paragraphs:
        if para_text.strip():
            if para_text.startswith("# "):
                doc.add_heading(para_text[2:], 1)
            elif para_text.startswith("## "):
                doc.add_heading(para_text[3:], 2)
            elif para_text.startswith("### "):
                doc.add_heading(para_text[4:], 3)
            else:
                doc.add_paragraph(para_text)

    doc.save(path)

    return {
        "success": True,
        "path": path,
        "paragraphs_added": len(paragraphs)
    }


def find_and_replace(path: str, find_text: str, replace_text: str) -> dict:
    """Find and replace text in a document."""
    if not Path(path).exists():
        return {"success": False, "error": f"File not found: {path}"}

    doc = Document(path)
    count = 0

    for para in doc.paragraphs:
        if find_text in para.text:
            # Simple replacement (loses some formatting)
            inline = para.runs
            for run in inline:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    count += 1

    doc.save(path)

    return {
        "success": True,
        "path": path,
        "replacements": count
    }


def get_document_structure(path: str) -> dict:
    """Get the structure/outline of a document."""
    if not Path(path).exists():
        return {"error": f"File not found: {path}"}

    doc = Document(path)
    structure = []

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""

        # Check if it's a heading
        if "Heading" in style_name:
            level = int(style_name.replace("Heading ", "")) if style_name != "Heading" else 1
            structure.append({
                "level": level,
                "text": para.text[:100],  # Truncate long headings
                "paragraph_index": i
            })

    return {
        "path": path,
        "total_paragraphs": len(doc.paragraphs),
        "structure": structure
    }


# Register tools with MCP server
@server.list_tools()
async def list_tools() -> list[Tool]:
    """List available docx tools."""
    return [
        Tool(
            name="docx_read",
            description="Read the content of a Word document as plain text",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute path to the .docx file"
                    }
                },
                "required": ["path"]
            }
        ),
        Tool(
            name="docx_read_structured",
            description="Read a Word document with structure info (paragraphs, tables, styles)",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute path to the .docx file"
                    }
                },
                "required": ["path"]
            }
        ),
        Tool(
            name="docx_write",
            description="Create a new Word document with content (supports markdown headings)",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute path for the new .docx file"
                    },
                    "content": {
                        "type": "string",
                        "description": "Content to write (paragraphs separated by double newlines, # for headings)"
                    },
                    "title": {
                        "type": "string",
                        "description": "Optional document title"
                    }
                },
                "required": ["path", "content"]
            }
        ),
        Tool(
            name="docx_append",
            description="Append content to an existing Word document",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute path to the .docx file"
                    },
                    "content": {
                        "type": "string",
                        "description": "Content to append"
                    }
                },
                "required": ["path", "content"]
            }
        ),
        Tool(
            name="docx_replace",
            description="Find and replace text in a Word document",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute path to the .docx file"
                    },
                    "find": {
                        "type": "string",
                        "description": "Text to find"
                    },
                    "replace": {
                        "type": "string",
                        "description": "Text to replace with"
                    }
                },
                "required": ["path", "find", "replace"]
            }
        ),
        Tool(
            name="docx_structure",
            description="Get the structure/outline of a Word document (headings and their positions)",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute path to the .docx file"
                    }
                },
                "required": ["path"]
            }
        )
    ]


@server.call_tool()
async def call_tool(name: str, arguments: dict[str, Any]) -> list[TextContent]:
    """Handle tool calls."""
    path = arguments.get("path", "")

    try:
        if name == "docx_read":
            if not Path(path).exists():
                return [TextContent(type="text", text=f"Error: File not found: {path}")]
            text = read_docx_plain(path)
            return [TextContent(type="text", text=text)]

        elif name == "docx_read_structured":
            if not Path(path).exists():
                return [TextContent(type="text", text=f"Error: File not found: {path}")]
            content = read_docx(path)
            return [TextContent(type="text", text=json.dumps(content, indent=2))]

        elif name == "docx_write":
            result = write_docx(
                path,
                arguments["content"],
                arguments.get("title")
            )
            return [TextContent(type="text", text=json.dumps(result, indent=2))]

        elif name == "docx_append":
            result = append_to_docx(path, arguments["content"])
            return [TextContent(type="text", text=json.dumps(result, indent=2))]

        elif name == "docx_replace":
            result = find_and_replace(
                path,
                arguments["find"],
                arguments["replace"]
            )
            return [TextContent(type="text", text=json.dumps(result, indent=2))]

        elif name == "docx_structure":
            structure = get_document_structure(path)
            return [TextContent(type="text", text=json.dumps(structure, indent=2))]

        else:
            return [TextContent(type="text", text=f"Unknown tool: {name}")]

    except Exception as e:
        return [TextContent(type="text", text=f"Error: {str(e)}")]


async def main():
    """Run the MCP server."""
    async with stdio_server() as (read_stream, write_stream):
        await server.run(read_stream, write_stream)


if __name__ == "__main__":
    asyncio.run(main())
